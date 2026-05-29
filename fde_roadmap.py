"""FDE转型路线图 — Word文档生成"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 样式设置 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.color.rgb = RGBColor(30, 40, 60)

# ── 封面 ──
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FDE 前沿部署工程师')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(25, 35, 55)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('个人转型路线图')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(80, 90, 120)

doc.add_paragraph('')
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('2026年5月30日\n').font.size = Pt(12)
info.add_run('基于OpenClaw三节点AI系统 + 六条业务线实战经验').font.size = Pt(11)

doc.add_page_break()

# ── 目录页 ──
doc.add_heading('目录', level=1)
toc_items = [
    '一、FDE是什么 & 为什么是你',
    '二、能力现状评估',
    '三、能力缺口分析',
    '四、三阶段转型路线',
    '五、第一阶段：基础搭建（1-30天）',
    '六、第二阶段：行业深耕（2-3月）',
    '七、第三阶段：品牌化输出（4-6月）',
    '八、FDE服务产品化方案',
    '九、客户获取策略',
    '十、关键里程碑 & 验收标准',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ── 第一章：FDE是什么 ──
doc.add_heading('一、FDE是什么 & 为什么是你', level=1)

doc.add_paragraph(
    'FDE（Forward Deployed Engineer，前沿部署工程师）是由 Palantir 首创、'
    '2025-2026年在硅谷AI行业爆发式增长的新型技术岗位。'
    '其核心是将软件工程能力与咨询顾问思维融合，直接驻扎客户现场，'
    '用AI技术解决真实的业务问题。'
)

doc.add_heading('行业数据', level=2)
p = doc.add_paragraph()
p.add_run('• ').font.bold = True
p.add_run('2025年4月至2026年4月，FDE岗位从643个飙升至5,330个（增长829%）\n')
p.add_run('• ').font.bold = True
p.add_run('YC孵化器超过100家AI公司正在招聘FDE（三年前为0）\n')
p.add_run('• ').font.bold = True
p.add_run('OpenAI组建The Deployment Company，募资超40亿美元，专做FDE部署\n')
p.add_run('• ').font.bold = True
p.add_run('上海2026年3月启动系统性FDE人才培育行动')

doc.add_heading('为什么这个方向适合你', level=2)
p = doc.add_paragraph(
    '你已经在不自知地做FDE的工作：用AI工具（Logo生成器、内容生成、'
    '数据分析）为客户解决具体问题，跨多个行业（贸易/餐饮/科技/心理），'
    '一人承担从需求沟通到最终交付的全流程。FDE不是从零开始，'
    '而是把你已经在做的事系统化、专业化、品牌化。'
)

doc.add_page_break()

# ── 第二章：能力现状评估 ──
doc.add_heading('二、能力现状评估', level=1)

doc.add_paragraph('对照FDE核心技能矩阵，评估现有能力和已有资产：')

table = doc.add_table(rows=11, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['能力维度', 'FDE要求', '你的现状', '评级']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)

rows_data = [
    ['Python开发', '精通，能快速原型开发', 'Logo生成器35KB/550+行，已实战', '⭐⭐⭐⭐'],
    ['AI/ML应用', 'LLM应用开发、RAG、Prompt Engineering', 'DeepSeek深度使用，多模型路由', '⭐⭐⭐'],
    ['系统部署', 'Docker/K8s/Linux', 'Windows环境为主，容器化未涉及', '⭐⭐'],
    ['前端/全栈', '快速搭建可交互界面', 'GitHub Pages静态站，无动态应用', '⭐⭐'],
    ['行业知识', '深耕1-2个行业', '贸易/跨境/餐饮/设计，广度>深度', '⭐⭐⭐'],
    ['客户沟通', '需求翻译、现场解决问题', '闲鱼/微信私域直接接单交付', '⭐⭐⭐⭐'],
    ['产品化思维', '沉淀可复用方案', 'Logo生成器已产品化(3档定价)', '⭐⭐⭐'],
    ['工具链', 'AI辅助开发工具', '三节点系统+OpenClaw+OKX监控', '⭐⭐⭐⭐'],
    ['英语能力', '技术文档读写', '可读英文技术文档', '⭐⭐⭐'],
    ['自驱力', '独立在客户现场工作', '一人跑6条业务线=极致自驱', '⭐⭐⭐⭐⭐'],
]
for i, row_data in enumerate(rows_data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val
        for p in table.rows[i+1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

doc.add_paragraph('')
doc.add_paragraph('核心优势：自驱力、客户沟通、工具搭建三项已超FDE入门标准。主要缺口在部署和全栈。')

doc.add_page_break()

# ── 第三章：能力缺口 ──
doc.add_heading('三、能力缺口分析', level=1)

table2 = doc.add_table(rows=6, cols=4)
table2.style = 'Light Grid Accent 1'
headers2 = ['缺口', '影响', '优先级', '预估投入']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)

gap_data = [
    ['Docker/K8s', '无法交付容器化方案给企业客户', '🟡 中', '2-3周'],
    ['前端框架(React/Vue)', '只能做静态页面，无法交付可交互应用', '🟡 中', '3-4周'],
    ['行业深度认证', '缺乏贸易/跨境领域权威背书', '🟢 低', '持续'],
    ['英文商务沟通', '无法接海外客户FDE项目', '🟢 低', '持续'],
    ['AI Agent框架', 'RAG/LangChain未系统学习', '🔴 高', '4-6周'],
]
for i, row_data in enumerate(gap_data):
    for j, val in enumerate(row_data):
        table2.rows[i+1].cells[j].text = val
        for p in table2.rows[i+1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

doc.add_page_break()

# ── 第四章：三阶段路线 ──
doc.add_heading('四、三阶段转型路线', level=1)

table3 = doc.add_table(rows=4, cols=4)
table3.style = 'Light Grid Accent 1'
headers3 = ['阶段', '时间', '核心目标', '关键产出']
for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)

stage_data = [
    ['第一阶段\n基础搭建', '1-30天\n(5/30-6/29)',
     '现有资产产品化\n补齐核心技能缺口',
     'FDE Portfolio文档\n3个可演示案例\nDocker基础认证'],
    ['第二阶段\n行业深耕', '2-3月\n(7月-8月)',
     '深耕贸易/跨境行业\n建立行业FDE方法论',
     '行业AI方案集\n2个深度客户案例\nLangChain实战项目'],
    ['第三阶段\n品牌化输出', '4-6月\n(9月-11月)',
     '对外输出FDE服务\n建立个人品牌',
     '付费FDE服务线\n行业分享/文章\n稳定客户管道'],
]
for i, row_data in enumerate(stage_data):
    for j, val in enumerate(row_data):
        table3.rows[i+1].cells[j].text = val
        for p in table3.rows[i+1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

doc.add_page_break()

# ── 第五章：第一阶段详细 ──
doc.add_heading('五、第一阶段：基础搭建（5/30-6/29）', level=1)

doc.add_heading('5.1 现有资产产品化（第1周）', level=2)

tasks_week1 = [
    ('构建FDE Portfolio文档', '将6条业务线重新包装为FDE案例格式：客户需求→AI方案→交付结果→可复用组件。每条业务线一页PPT/一页文档。'),
    ('Logo生成器升级', '补全说明文档、添加API接口、录制Demo视频（屏幕录制+配音）。目标：一个链接就能展示完整能力。'),
    ('建立方案模板库', 'Logo方案模板、简历优化模板、商业计划书模板——每个模板含需求问卷+AI处理流程+交付物范例。'),
]
for title, desc in tasks_week1:
    p = doc.add_paragraph()
    p.add_run(f'■ {title}：').font.bold = True
    p.add_run(desc)

doc.add_heading('5.2 核心技术补缺（第2-3周）', level=2)

tasks_week2 = [
    ('Docker基础', '学习容器化概念，将Logo生成器打包成Docker镜像。目标：docker run 一键启动。在线课程推荐：Docker官方入门教程（免费）。'),
    ('GitHub完善', '将Logo生成器、记账工具等推送到GitHub公开仓库。写好README（中英双语）。这是FDE的"代码名片"。'),
    ('Streamlit/Gradio入门', '用Python快速搭建Web UI。将Logo生成器从命令行升级为网页交互界面。学习时间：2-3天。'),
]
for title, desc in tasks_week2:
    p = doc.add_paragraph()
    p.add_run(f'■ {title}：').font.bold = True
    p.add_run(desc)

doc.add_heading('5.3 第一个FDE案例深度打磨（第4周）', level=2)

p = doc.add_paragraph()
p.add_run('选择"Logo设计"作为第一个完整FDE案例，因为它最成熟：\n').font.bold = True
p.add_run('• 写一份完整的Case Study：客户背景→原始需求→技术选型→迭代过程（v1→v2）→最终交付\n'
          '• 量化结果：从零设计基础到32版/品牌方案，交付时间缩短90%\n'
          '• 抽象出可复用方法论：AI辅助视觉设计的三步法\n'
          '• 发布到知乎/掘金/小红书（长文版）')

doc.add_page_break()

# ── 第六章：第二阶段 ──
doc.add_heading('六、第二阶段：行业深耕（7月-8月）', level=1)

doc.add_heading('6.1 选定深耕行业：大宗贸易 + 跨境电商', level=2)

p = doc.add_paragraph()
p.add_run('选择理由：\n').font.bold = True
p.add_run('• 你已有温氏饲料方案（14章53表）的行业深度经验\n'
          '• 大宗贸易信息化程度低，AI改造空间极大\n'
          '• 跨境电商是AI应用最活跃的赛道之一\n'
          '• 两个行业互补：一个是传统toB，一个是数字化toC')

doc.add_heading('6.2 行业FDE方案构建', level=2)

tasks_ind = [
    ('贸易行业AI方案集', '合同条款AI审查、报价单自动生成、物流路线AI优化、供应商评估AI系统、关税合规AI检查。每项先出MVP再迭代。'),
    ('跨境电商AI方案集', '选品AI分析、Listing优化、竞品监控、库存预测、客服自动回复。已有亚马逊店铺可做实测。'),
    ('LangChain系统学习', '完成LangChain官方教程（2周）。用LangChain搭建一个RAG应用：将你的行业知识库（贸易文档/跨境资料）做成可查询的AI助手。'),
    ('第一个企业级Demo', '选一个最小可行场景（如"合同AI审查"），用Streamlit+LangChain做一个完整Demo，录制成5分钟演示视频。'),
]
for title, desc in tasks_ind:
    p = doc.add_paragraph()
    p.add_run(f'■ {title}：').font.bold = True
    p.add_run(desc)

doc.add_heading('6.3 行业人脉与背书', level=2)
p = doc.add_paragraph(
    '利用已有资源——凌丰集团/万事泰集团/建材圈6位老板——做一件事：'
    '免费为他们做一次AI能力展示（如合同AI审查Demo），换取案例授权和推荐。'
    'FDE的核心信任来源于"做过"，而非"学过"。'
)

doc.add_page_break()

# ── 第七章：第三阶段 ──
doc.add_heading('七、第三阶段：品牌化输出（9月-11月）', level=1)

doc.add_heading('7.1 建立个人FDE品牌', level=2)

p = doc.add_paragraph()
p.add_run('线上：\n').font.bold = True
p.add_run('• 知乎/掘金：发布FDE系列技术文章（月均2篇）\n'
          '• 小红书号1继续运营，增加FDE/Prompt Engineering内容\n'
          '• GitHub保持活跃，每个项目有完整的README + Demo\n'
          '• 考虑录制B站教学视频：零基础学AI部署\n\n')
p.add_run('线下：\n').font.bold = True
p.add_run('• 通过贸易圈子介绍AI改造案例\n'
          '• 参与本地AI/创业社群活动\n'
          '• 目标：3个月内积累5个真实案例 + 3个行业推荐')

doc.add_heading('7.2 FDE服务商业化', level=2)

table4 = doc.add_table(rows=4, cols=4)
table4.style = 'Light Grid Accent 1'
for i, h in enumerate(['服务档位', '内容', '定价', '对标']):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)

pricing = [
    ['AI诊断', '行业AI机会评估\n1天驻场/远程调研\n输出AI改造建议书', '¥1,999-4,999', '咨询入门'],
    ['AI部署', '选定场景AI方案落地\nDocker部署+培训\n3个月维护', '¥15,000-50,000', 'FDE标准'],
    ['AI合伙人', '长期驻场/陪跑\n多个场景逐一部署\n按效果分成', '股权/分成', '高级FDE'],
]
for i, row_data in enumerate(pricing):
    for j, val in enumerate(row_data):
        table4.rows[i+1].cells[j].text = val
        for p in table4.rows[i+1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

doc.add_page_break()

# ── 第八章：服务产品化 ──
doc.add_heading('八、FDE服务产品化方案', level=1)

doc.add_paragraph(
    '将每一条业务线重新定义为FDE解决方案，而非零散服务：'
)

table5 = doc.add_table(rows=7, cols=5)
table5.style = 'Light Grid Accent 1'
for i, h in enumerate(['原业务线', 'FDE重新定位', '核心AI能力', '目标客户', '可复用资产']):
    cell = table5.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(8)

biz_data = [
    ['Logo设计', 'AI视觉设计部署', 'Python PIL生成\n多风格自动适配', '中小企业\n初创公司', 'Logo生成器\n32版/品牌方案'],
    ['简历/演讲稿', 'AI内容生成部署', 'Prompt Engineering\n模板系统', '求职者\n企业管理者', '内容模板库\n一键生成工具'],
    ['商业计划书', 'AI商业分析部署', '结构化分析\n数据可视化', '创业者\n融资团队', '14章53表模板\n行业数据库'],
    ['记账工具', 'AI财税自动化', '数据抓取\n自动分类汇总', '小微企业\n个体户', '订阅账单汇总\n记账小助手'],
    ['合同审查', 'AI合规审查部署', 'LangChain RAG\n条款对比引擎', '贸易公司\n法律顾问', '合同模板库\n风险点标签系统'],
    ['跨境电商', 'AI电商运营部署', '选品分析\nListing优化\n竞品监控', '亚马逊卖家\n独立站运营', '运营Dashboard\n自动化脚本集'],
]
for i, row_data in enumerate(biz_data):
    for j, val in enumerate(row_data):
        table5.rows[i+1].cells[j].text = val
        for p in table5.rows[i+1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

doc.add_page_break()

# ── 第九章：客户获取 ──
doc.add_heading('九、客户获取策略', level=1)

doc.add_heading('9.1 三级客户管道', level=2)

p = doc.add_paragraph()
p.add_run('第一层：现有关系转化（立即启动）\n').font.bold = True
p.add_run('• 凌丰集团/万事泰集团/建材圈6位老板 → 免费AI诊断换取案例\n'
          '• 粤盐贸易（正在沟通）→ 贸易行业AI方案试点\n'
          '• 家电店朋友（做直播）→ 直播数据分析AI方案\n\n')
p.add_run('第二层：平台获客（第2-3个月启动）\n').font.bold = True
p.add_run('• 闲鱼：增加"企业AI部署"商品\n'
          '• 知乎：发布FDE案例文章引流\n'
          '• 小红书：展示AI改造效果\n'
          '• LinkedIn：英文Profile建FDE标签\n\n')
p.add_run('第三层：口碑裂变（第4-6个月）\n').font.bold = True
p.add_run('• 每个交付项目要求客户推荐1个潜在客户\n'
          '• 参与行业论坛/社群分享\n'
          '• 与本地商会/产业园区合作')

doc.add_heading('9.2 获客话术框架', level=2)

p = doc.add_paragraph(
    '不说"我是做AI的"，而是说"我可以帮你的公司用AI省掉X成本/提Y效率"。'
    '每个客户接触前，先做3件事：了解对方行业痛点 → 准备一个针对性Demo → 量化为具体数字。'
)

doc.add_page_break()

# ── 第十章：里程碑 ──
doc.add_heading('十、关键里程碑 & 验收标准', level=1)

table6 = doc.add_table(rows=8, cols=4)
table6.style = 'Light Grid Accent 1'
for i, h in enumerate(['时间节点', '里程碑', '验收标准', '若不达标']):
    cell = table6.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)

milestones = [
    ['第1周末\n(6/6)', 'Portfolio v1完成', '6条业务线均以FDE格式呈现\nLogo生成器Docker化', '延期不超过1周'],
    ['第3周末\n(6/20)', '核心技能补缺', 'Docker基础认证/课程完成\nGitHub 3个Repo有完整README', '继续学习不跳过'],
    ['第4周末\n(6/29)', '第一个FDE案例发布', 'Logo设计Case Study完成\n知乎/掘金发布，阅读量>500', '优化后7月初发布'],
    ['第6周末\n(7/13)', 'Streamlit Demo上线', 'Logo生成器Web界面可访问\n录5分钟Demo视频', '先做截图版'],
    ['第8周末\n(7/27)', 'LangChain项目完成', '合同AI审查RAG原型可运行\n能处理3种以上合同类型', '缩小到1种类型'],
    ['第12周末\n(8/24)', '行业方案集完成', '贸易/跨境各3个AI方案MVP\n至少1个方案有真实客户使用', '聚焦1个行业'],
    ['第24周末\n(11/15)', 'FDE服务商业化', '月均1个付费客户\n个人品牌月均曝光>5000', '调整定价或渠道'],
]
for i, row_data in enumerate(milestones):
    for j, val in enumerate(row_data):
        table6.rows[i+1].cells[j].text = val
        for p in table6.rows[i+1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— 路线图版本 v1.0 · 2026年5月30日 · 基于OpenClaw三节点系统 ——')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(130, 130, 150)

# ── 保存 ──
out_dir = r"C:\Users\YH2\Desktop\CC文档"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "FDE转型路线图_v1.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
