"""FDE Portfolio文档生成 — 6条业务线转FDE案例格式"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.color.rgb = RGBColor(30, 40, 60)

# ── 封面 ──
for _ in range(4):
    doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FDE Portfolio')
run.font.size = Pt(30)
run.font.bold = True
run.font.color.rgb = RGBColor(25, 35, 55)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('前沿部署工程师 · 案例集')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 90, 120)

doc.add_paragraph('')
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.add_run('6条业务线 × AI落地方案\n').font.size = Pt(12)
sub2.add_run('OpenClaw三节点系统 · 2026年5月').font.size = Pt(10)

doc.add_page_break()

# ── 案例模板函数 ──
def add_case(doc, num, title, client_need, ai_solution, delivery_result, reusable_assets, biz_line):
    doc.add_heading(f'案例 #{num}：{title}', level=2)
    
    # 元信息
    p = doc.add_paragraph()
    p.add_run(f'业务线：{biz_line}').font.size = Pt(9)
    p.add_run(f'    状态：已交付 / 运营中').font.size = Pt(9)
    
    # 表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    labels = ['🎯 客户需求', '🤖 AI方案', '📊 交付结果', '♻️ 可复用资产', '💡 FDE价值']
    values = [client_need, ai_solution, delivery_result, reusable_assets,
              f'展示了一名FDE从{title.split("：")[0] if "：" in title else title}需求出发，'
              f'用AI工具快速搭建完整交付链路的能力。从问题理解→技术选型→原型开发→客户交付，全流程独立完成。']
    
    for i, (label, val) in enumerate(zip(labels, values)):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = val
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph('')

# ── 案例1：Logo设计 ──
add_case(doc, 1, 'Logo设计：AI视觉设计部署',
    '中小企业/初创公司需要专业Logo，但预算有限（几百到几千不等），'
    '传统渠道找设计师周期长（3-7天），沟通成本高，风格选择少（1-3版）。'
    '客户需要的不是"AI工具"，而是"一个能直接用的Logo"。',
    
    '搭建Python/Pillow Logo自动生成引擎（550+行代码），整合8种视觉调性系统'
    '（modern/elegance/bold/playful/tech/vintage/minimal/nature），'
    '14个行业自动配色和字体适配。开发Streamlit Web界面，'
    '客户零门槛使用。三档定价：体验版¥69(2方向)→标准版¥129(4方向)→进阶版¥299(8方向)。',
    
    '• 设计成本降幅：¥500-5000 → ¥69-299（降90%+）\n'
    '• 交付速度：3-7天 → 3秒\n'
    '• 方案多样性：1-3版 → 最多32版/品牌\n'
    '• 已产出真实案例：KH投资、谊璜贸易、粤盐等多个品牌',
    
    '• Logo生成器核心引擎（generate.py，35KB）\n'
    '• Streamlit Web界面（fde_studio.py）\n'
    '• Docker容器化配置\n'
    '• 方案对比拼图工具\n'
    '• GitHub公开仓库（fde-logo-studio）',
    'Logo设计')

doc.add_page_break()

# ── 案例2 ──
add_case(doc, 2, '简历优化：AI内容生成部署',
    '求职者在闲鱼/招聘平台寻找简历优化服务，传统人工优化收费¥200-800/份，'
    '交付需1-3天。客户痛点：不知道怎么写才能过ATS筛选、突出优势。'
    '需要的不只是改错别字，而是结构重组+亮点提炼。',
    
    '基于DeepSeek大模型，构建简历优化Prompt工程体系。建立职位-简历匹配分析框架，'
    '支持中英文双语优化。三档定价：基础优化¥49、标准改版¥99、精修重塑¥199。'
    '输出包含：优化版简历+修改说明+ATS关键词匹配报告。',
    
    '• 闲鱼已上架在售\n'
    '• 从接单到交付全流程1人完成\n'
    '• 客户复购率（演讲稿等衍生需求）',
    
    '• 简历优化Prompt模板库\n'
    '• 行业JD关键词数据库\n'
    '• ATS兼容性检查清单',
    '闲鱼服务')

# ── 案例3 ──
add_case(doc, 3, '商业计划书：AI商业分析部署',
    '创业者/融资团队需要专业商业计划书，但不知道怎么写、写什么。'
    '传统代写费用¥2000-8000/份，周期2-4周。'
    '客户需要的是：行业数据支撑+财务预测模型+投资人视角。',
    
    '构建14章53表的标准化商业计划书模板体系（基于温氏饲料方案实战经验）。'
    '整合AI生成+数据分析+可视化输出。典型报告包含：行业分析/竞争格局/财务预测/风险评估。'
    '二档定价：标准版¥399、深度版¥699。',
    
    '• 温氏饲料方案：14章53表Word完整报告\n'
    '• 年用量数据：玉米1650-2100万吨/豆粕450-700万吨\n'
    '• 进口成本比东北到港省100-150元/吨\n'
    '• 已作为大宗贸易切入方案实际使用',
    
    '• 14章53表标准模板\n'
    '• 行业数据库（农业/贸易）\n'
    '• 财务预测模型框架',
    '代写服务')

doc.add_page_break()

# ── 案例4 ──
add_case(doc, 4, '记账工具：AI财税自动化部署',
    '小微企业/个体户/自由职业者需要记账但嫌专业软件复杂且贵。'
    '订阅制服务（视频/云盘/会员等）账单分散在各平台，汇总管理困难。'
    '客户需要的是"傻瓜式"记账，而非学习一套财务系统。',
    
    '开发订阅账单汇总工具，自动抓取各平台账单数据，分类汇总后生成可视化报表。'
    '部署在GitHub Pages静态站（yhmy2025.github.io），零运维成本。'
    '免费工具引流→弹窗引导加微信/搜闲鱼→转化简历/Logo等高客单价服务。',
    
    '• 已部署至yhmy2025.github.io\n'
    '• 配套ToolChecker.ps1监控脚本\n'
    '• 形成"免费工具→私域→付费服务"三级漏斗',
    
    '• 记账小助手Web工具\n'
    '• 订阅账单汇总脚本\n'
    '• ToolChecker监控脚本\n'
    '• 私域引流漏斗模型',
    '私域工具')

# ── 案例5 ──
add_case(doc, 5, '小红书运营：AI内容部署',
    '双号运营需要每日产出高质量内容，但个人精力有限。'
    '号1"YH的副业日记"（AI副业记录）和号2"心理树洞号"（深夜嘴替）内容方向不同，'
    '需要定制化内容策略+持续输出能力。',
    
    '基于DeepSeek大模型+OpenClaw框架，搭建自动化内容管线：'
    '选题库（35个选题按类型分类）→AI写稿→风格自检→配图生成→发布记录。'
    '号2"嘴替"定位经A/B测试验证（矛盾型>场景描写>教程型）。'
    '目标：搭建定时推送系统（每晚19:30自动选题→写稿→推送微信bot）。',
    
    '• 号1已发布7篇\n'
    '• 号2已发布4篇，验证了嘴替内容模型\n'
    '• 建立发布记录+选题库（35个选题）\n'
    '• Pillow排版海报生成工具',
    
    '• 35个选题库\n'
    '• 内容风格自检流程\n'
    '• 封面图生成工具\n'
    '• 发布记录追踪表',
    '自媒体运营')

doc.add_page_break()

# ── 案例6 ──
add_case(doc, 6, '加密货币交易：AI交易监控部署',
    '加密货币7×24小时交易，无法人工盯盘。需要自动监控仓位、'
    '聪明钱信号、多空比等指标，在特定触发条件下推送通知。'
    '核心需求：减少噪音推送、只在关键时刻提醒。',
    
    '基于OKX API + OpenClaw cron/heartbeat框架，搭建多层监控系统：\n'
    '• 心跳层：每30分钟检查仓位+聪明钱+价格，仅在触发条件时推送\n'
    '• 定时层：每日3次市场快照，自动推送到微信Bot\n'
    '• 触发条件：价格异动>3%/聪明钱翻转/仓位逼近TP-SL/极端市场信号\n'
    '• 子代理层：每周自动生成市场分析报告',
    
    '• 日均自动监控48次，无触发不打扰\n'
    '• 多空判断结合聪明钱信号+RSI+OI\n'
    '• 微信Bot实时推送\n'
    '• 全自动运行，人工介入仅需下单',
    
    '• OKX API集成配置\n'
    '• 监控触发规则模板\n'
    '• 每日报告自动生成脚本\n'
    '• 风险控制参数体系',
    '加密货币')

# ── 总结页 ──
doc.add_page_break()
doc.add_heading('总结：FDE能力矩阵', level=1)

table = doc.add_table(rows=8, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['FDE核心能力', '案例证据', '水平', '下一步']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)

mat = [
    ['需求翻译\n技术↔业务', '案例1/2/3：客户需求→技术方案', '⭐⭐⭐⭐', '增加行业术语库'],
    ['快速原型开发', '案例1：Logo引擎35KB/550行\n案例4：记账工具', '⭐⭐⭐⭐', '补充前端框架'],
    ['客户现场交付', '案例1/2：闲鱼/微信直接交付', '⭐⭐⭐⭐', '拓展企业客户'],
    ['产品化思维', '案例1：三档定价+打包下载\n案例4：免费→付费漏斗', '⭐⭐⭐', '完善定价体系'],
    ['AI工具集成', '全案例：DeepSeek+OKX+Streamlit', '⭐⭐⭐⭐', '学习LangChain'],
    ['部署与运维', '案例4：GitHub Pages\n案例6：cron心跳', '⭐⭐⭐', 'Docker+K8s'],
    ['数据驱动决策', '案例5：A/B测试内容模型\n案例6：多指标监控', '⭐⭐⭐⭐', '建立Dashboard'],
]
for i, row in enumerate(mat):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val
        for p in table.rows[i+1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('🏷️ FDE定位：').font.bold = True
p.add_run('AI驱动的多行业解决方案部署专家，覆盖视觉设计、内容生成、数据分析、自动化监控四大领域。')

p2 = doc.add_paragraph()
p2.add_run('🎯 目标客户：').font.bold = True
p2.add_run('中小企业（30-200人）、创业公司、个体经营者，有信息化需求但无技术团队。')

p3 = doc.add_paragraph()
p3.add_run('💰 服务模式：').font.bold = True
p3.add_run('轻AI诊断（¥1,999-4,999）→ AI部署（¥15,000-50,000）→ AI合伙人（分成/股权）。')

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— FDE Portfolio v1.0 · 2026-05-30 · OpenClaw三节点系统 ——')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(130, 130, 150)

# ── 保存 ──
out_dir = r"C:\Users\YH2\Desktop\CC文档"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "FDE_Portfolio_案例集_v1.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
